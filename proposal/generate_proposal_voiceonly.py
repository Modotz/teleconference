# -*- coding: utf-8 -*-
"""
Generator dokumen Word: Proposal Penjualan SDK Voice Call (TANPA Video Meeting).
Jalankan:  python generate_proposal_voiceonly.py
Output  :  Proposal-SDK-VoiceCall-Only.docx
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0x25, 0x63, 0xEB)
GREY = RGBColor(0x55, 0x55, 0x55)
HEADER_FILL = "1F3A5F"
ALT_FILL = "EEF2F7"

doc = Document()

section = doc.sections[0]
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.2)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def h1(text):
    p = doc.add_heading(level=1)
    run = p.add_run(text)
    run.font.color.rgb = NAVY
    run.font.size = Pt(16)
    return p


def h2(text):
    p = doc.add_heading(level=2)
    run = p.add_run(text)
    run.font.color.rgb = ACCENT
    run.font.size = Pt(13)
    return p


def para(text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align is not None:
        p.alignment = align
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def mono(lines):
    p = doc.add_paragraph()
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        shade_cell(hdr[i], HEADER_FILL)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            run = cells[ci].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            if ri % 2 == 1:
                shade_cell(cells[ci], ALT_FILL)
    if widths:
        for row in t.rows:
            for ci, w in enumerate(widths):
                row.cells[ci].width = Cm(w)
    return t


# ===========================================================================
# COVER
# ===========================================================================
for _ in range(3):
    doc.add_paragraph()

para("PROPOSAL PENAWARAN", bold=True, size=14, color=ACCENT,
     align=WD_ALIGN_PARAGRAPH.CENTER)
para("Solusi SDK Voice Call", bold=True, size=28, color=NAVY,
     align=WD_ALIGN_PARAGRAPH.CENTER)
para("SDK Panggilan Suara Real-Time untuk Integrasi Aplikasi — "
     "Kapasitas 500 Pengguna", size=12, color=GREY,
     align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(3):
    doc.add_paragraph()

table(
    ["Keterangan", "Detail"],
    [
        ["Diajukan untuk", "[Nama Klien / Instansi]"],
        ["Disusun oleh", "[Nama Perusahaan Anda]"],
        ["Nomor Dokumen", "PROP/VC/2026/002"],
        ["Tanggal", "Mei 2026"],
        ["Versi Dokumen", "1.0"],
        ["Lingkup Penawaran", "SDK Voice Call saja (tanpa Video Meeting)"],
        ["Masa Berlaku Penawaran", "30 hari kalender sejak tanggal terbit"],
    ],
    widths=[6, 10],
)

for _ in range(2):
    doc.add_paragraph()
para("Dokumen ini bersifat rahasia dan ditujukan khusus untuk pihak penerima. "
     "Dilarang menggandakan atau menyebarluaskan tanpa izin tertulis.",
     italic=True, size=9, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ===========================================================================
# 1. RINGKASAN EKSEKUTIF
# ===========================================================================
h1("1. Ringkasan Eksekutif")
para(
    "Kami menawarkan SDK Voice Call — sebuah Software Development Kit untuk "
    "menanamkan fitur panggilan suara real-time ke dalam aplikasi mana pun, "
    "termasuk aplikasi berbasis Next.js / React. Solusi ini ditujukan untuk "
    "melayani hingga 500 pengguna dan di-deploy secara mandiri (self-hosted) "
    "pada server milik klien maupun cloud."
)
para(
    "Dengan SDK ini, tim pengembang aplikasi pihak ketiga dapat menambahkan "
    "panggilan suara 1-lawan-1 maupun grup hanya dengan beberapa baris kode — "
    "tanpa perlu membangun infrastruktur WebRTC dari nol. Peserta panggilan "
    "tidak perlu membuat akun; identitas cukup melalui token akses berumur "
    "pendek yang diterbitkan oleh server aplikasi klien."
)
para("Catatan lingkup:", bold=True)
bullet("Penawaran ini KHUSUS untuk SDK Voice Call (panggilan suara).")
bullet("Fitur Video Meeting, berbagi layar, dan perekaman TIDAK termasuk; "
       "tersedia pada paket terpisah.")

# ===========================================================================
# 2. LATAR BELAKANG & TUJUAN
# ===========================================================================
h1("2. Latar Belakang & Tujuan")
para(
    "Membangun fitur panggilan suara berbasis WebRTC dari awal membutuhkan "
    "keahlian khusus (media server, signaling, NAT traversal) dan waktu "
    "pengembangan yang panjang. SDK Voice Call hadir sebagai komponen siap "
    "pakai sehingga organisasi dapat fokus pada produk inti mereka."
)
para("Tujuan penyediaan solusi ini:", bold=True)
bullet("menyediakan SDK panggilan suara yang ringan dan mudah diintegrasikan;")
bullet("memungkinkan aplikasi pihak ketiga menambahkan voice call dengan cepat;")
bullet("menjamin keamanan komunikasi melalui enkripsi standar WebRTC;")
bullet("memberikan kendali penuh atas infrastruktur dan data kepada klien.")

# ===========================================================================
# 3. FITUR SDK VOICE CALL
# ===========================================================================
h1("3. Fitur SDK Voice Call")

h2("3.1. Kemampuan Panggilan Suara")
for f in [
    "Panggilan suara 1-lawan-1 dan panggilan grup (multi-peserta)",
    "Arsitektur SFU — efisien untuk panggilan dengan banyak peserta",
    "Audio jernih dengan codec Opus",
    "Kontrol mute / unmute",
    "Deteksi pembicara aktif (active speaker detection)",
    "Notifikasi panggilan masuk (incoming call)",
    "Latensi rendah, kualitas adaptif terhadap kondisi jaringan",
]:
    bullet(f)

h2("3.2. SDK & Cara Integrasi")
for f in [
    "SDK JavaScript / TypeScript: paket @teleconf/voicecall-sdk",
    "Komponen siap pakai untuk Next.js / React: <VoiceCallRoom>",
    "React Hook headless (useVoiceCall) untuk membangun UI sendiri",
    "Kelas VoiceCallClient untuk kontrol penuh tingkat lanjut",
    "Pemutaran audio peserta lain otomatis ditangani oleh SDK",
    "Integrasi ringkas: terbitkan token di server, sematkan di sisi klien",
]:
    bullet(f)

h2("3.3. REST API & Dokumentasi")
for f in [
    "REST API publik (v1): buat panggilan, terbitkan token, status, akhiri",
    "Autentikasi melalui API Key (dipakai di sisi server aplikasi klien)",
    "Access token berumur pendek, terbatas pada satu panggilan, per peserta",
    "Peserta sebagai tamu — tanpa perlu membuat akun",
    "Dokumentasi interaktif OpenAPI 3.0 / Swagger UI",
]:
    bullet(f)

h2("3.4. Keamanan")
for f in [
    "API Key disimpan dalam bentuk hash (SHA-256)",
    "Access token berumur pendek dan ber-scope per panggilan",
    "Media audio terenkripsi (DTLS-SRTP, standar WebRTC)",
    "Seluruh komunikasi melalui HTTPS / WSS",
]:
    bullet(f)

doc.add_page_break()

# ===========================================================================
# 4. ARSITEKTUR & TOPOLOGI
# ===========================================================================
h1("4. Arsitektur & Topologi Sistem")
para(
    "Sistem menggunakan arsitektur SFU (Selective Forwarding Unit). Setiap "
    "peserta mengirim satu aliran audio ke server, lalu server meneruskannya "
    "ke peserta lain. Pendekatan ini efisien dan ringan, khususnya untuk "
    "panggilan suara dengan banyak peserta."
)

h2("4.1. Diagram Topologi")
mono([
    "        +------------- APLIKASI PIHAK KETIGA -------------+",
    "        |                                                 |",
    "   [ Browser/Web ]   [ Mobile Web ]   [ Aplikasi Lain ]   ",
    "   [ + Voice SDK ]   [ + Voice SDK]   [ + Voice SDK   ]   ",
    "        |                  |                  |          ",
    "        +--------+---------+---------+--------+           ",
    "                 |  HTTPS / WSS / WebRTC (UDP+TCP)        ",
    "                 v                                       ",
    "        ==============  INTERNET  =====================   ",
    "                 |                                       ",
    "                 v                                       ",
    "   +-------------------------------------------------+   ",
    "   |          SERVER (Cloud / On-Premise)            |   ",
    "   |                                                 |   ",
    "   |   [ Nginx ]  reverse proxy + TLS (HTTPS/WSS)     |   ",
    "   |       |                                         |   ",
    "   |       +--> [ Backend  ] Node.js + Socket.io      |   ",
    "   |       |    [ Mediasoup] WebRTC SFU (audio)       |   ",
    "   |       +--> [ REST API v1 ] untuk SDK             |   ",
    "   |                    |                            |   ",
    "   |                    v                            |   ",
    "   |            [ PostgreSQL ] basis data             |   ",
    "   +-------------------------------------------------+   ",
])
para("Antarmuka pengguna (UI) disediakan oleh aplikasi pihak ketiga; SDK hanya "
     "menangani logika panggilan dan media. Signaling memakai WebSocket, "
     "media audio mengalir via WebRTC.", italic=True, size=9, color=GREY)

h2("4.2. Komponen Sistem")
table(
    ["Lapisan", "Komponen", "Fungsi"],
    [
        ["Klien", "Aplikasi Pihak Ketiga + Voice SDK",
         "Menjalankan SDK; UI milik aplikasi klien"],
        ["Gateway", "Nginx", "Reverse proxy, terminasi TLS"],
        ["Aplikasi", "Backend Node.js", "Signaling, REST API, autentikasi"],
        ["Media", "Mediasoup (SFU)", "Penerusan aliran audio WebRTC"],
        ["Data", "PostgreSQL", "Data API client, panggilan, audit"],
    ],
    widths=[2.6, 5.4, 8.0],
)

h2("4.3. Alur Integrasi SDK")
mono([
    "  1. Backend Partner --X-Api-Key--> POST /api/v1/calls           -> callId",
    "  2. Backend Partner --X-Api-Key--> POST /api/v1/calls/:id/token -> token",
    "  3. Browser Partner --accessToken-> SDK Voice Call --WebRTC--> Mediasoup",
])
para("API Key hanya dipakai di sisi server aplikasi klien. Browser hanya "
     "menerima access token berumur pendek untuk satu panggilan tertentu.",
     italic=True, size=9, color=GREY)

doc.add_page_break()

# ===========================================================================
# 5. TEKNOLOGI YANG DIGUNAKAN
# ===========================================================================
h1("5. Teknologi yang Digunakan")
para("SDK dan layanan pendukung dibangun dengan teknologi modern, open "
     "standard, dan teruji di lingkungan produksi.")
table(
    ["Kategori", "Teknologi", "Keterangan"],
    [
        ["SDK", "TypeScript, mediasoup-client, tsup",
         "Paket terdistribusi untuk aplikasi pihak ketiga"],
        ["Kompatibilitas SDK", "Next.js, React (>= 18)",
         "Komponen siap pakai + React Hook"],
        ["Backend", "Node.js 20+, Express, Socket.io",
         "REST API & signaling real-time"],
        ["Media Server", "Mediasoup (WebRTC SFU)",
         "Penerusan media audio efisien & skalabel"],
        ["Basis Data", "PostgreSQL 16 + Prisma ORM",
         "Penyimpanan API client & data panggilan"],
        ["Real-Time", "WebRTC, WebSocket",
         "Audio & signaling latensi rendah"],
        ["Audio Codec", "Opus", "Kualitas suara tinggi, hemat bandwidth"],
        ["Dokumentasi API", "OpenAPI 3.0, Swagger UI",
         "Referensi REST API interaktif"],
        ["Deployment", "Docker, Nginx, Let's Encrypt",
         "Kontainerisasi, reverse proxy, sertifikat TLS"],
        ["Keamanan", "API Key (hash), JWT, DTLS-SRTP",
         "Autentikasi & enkripsi media"],
    ],
    widths=[3.2, 5.8, 7.0],
)

# ===========================================================================
# 6. SPESIFIKASI SERVER
# ===========================================================================
h1("6. Spesifikasi Server")
para(
    "Panggilan suara (voice-only) jauh lebih ringan dibanding video. Codec "
    "Opus hanya memakai sekitar 40-50 kbps per aliran, sehingga kebutuhan CPU "
    "dan bandwidth server relatif kecil. Spesifikasi berikut diperhitungkan "
    "untuk 500 pengguna terdaftar dengan asumsi puncak penggunaan bersamaan "
    "(concurrent) sekitar 100-150 pengguna aktif dalam panggilan."
)

h2("6.1. Server Produksi (Direkomendasikan)")
table(
    ["Komponen", "Spesifikasi Minimum", "Spesifikasi Direkomendasikan"],
    [
        ["CPU", "2 vCPU", "4 vCPU (Intel Xeon / AMD EPYC generasi baru)"],
        ["RAM", "4 GB", "8 GB"],
        ["Penyimpanan", "40 GB SSD", "80 GB SSD"],
        ["Jaringan", "100 Mbps", "1 Gbps"],
        ["Sistem Operasi", "Ubuntu Server 22.04 LTS", "Ubuntu Server 22.04 LTS"],
        ["IP Publik", "1 IP statis", "1 IP statis"],
        ["Port", "80, 443, 40000-40100 (UDP+TCP)",
         "80, 443, 40000-40100 (UDP+TCP)"],
    ],
    widths=[3.2, 5.4, 7.4],
)

h2("6.2. Catatan Kapasitas & Skalabilitas")
bullet("Voice-only sangat efisien — satu server 4 vCPU mampu melayani "
       "ratusan peserta panggilan bersamaan dengan nyaman.")
bullet("Bandwidth jauh lebih kecil dari video; 200 peserta audio aktif "
       "umumnya masih di bawah 100 Mbps total.")
bullet("Mediasoup berjalan satu worker per inti CPU; menambah inti CPU "
       "meningkatkan kapasitas secara linear.")
bullet("Untuk pertumbuhan di atas 500 pengguna, sistem dapat di-scale "
       "horizontal dengan menambah node media.")
bullet("Database PostgreSQL dapat ditempatkan pada server yang sama tanpa "
       "membebani sistem secara berarti.")

h2("6.3. Lingkungan Pengujian (Opsional)")
table(
    ["Komponen", "Spesifikasi"],
    [
        ["CPU", "2 vCPU"],
        ["RAM", "4 GB"],
        ["Penyimpanan", "30 GB SSD"],
        ["Kegunaan", "Staging / UAT / pelatihan"],
    ],
    widths=[5.0, 11.0],
)

doc.add_page_break()

# ===========================================================================
# 7. ESTIMASI BIAYA
# ===========================================================================
h1("7. Estimasi Biaya")
para("Tersedia dua skema kepemilikan. Klien dapat memilih sesuai kebutuhan "
     "anggaran dan model operasional.", bold=True)
para("Seluruh harga dalam Rupiah (IDR), belum termasuk PPN. Angka bersifat "
     "estimasi dan dapat disesuaikan berdasarkan kesepakatan akhir serta "
     "ruang lingkup final.", italic=True, size=9, color=GREY)

h2("Opsi A — Beli Putus (Perpetual License)")
table(
    ["No", "Komponen", "Deskripsi", "Estimasi Biaya (IDR)"],
    [
        ["1", "Lisensi SDK Voice Call",
         "SDK Voice Call, kapasitas 500 user, lisensi perpetual",
         "45.000.000"],
        ["2", "Implementasi & Deployment",
         "Instalasi, konfigurasi server, hardening, go-live",
         "10.000.000"],
        ["3", "Integrasi SDK",
         "Pendampingan integrasi SDK ke aplikasi pihak ketiga",
         "10.000.000"],
        ["4", "Pelatihan & Dokumentasi",
         "Pelatihan developer, serah-terima dokumen",
         "4.000.000"],
        ["", "TOTAL BIAYA AWAL (One-Time)", "", "69.000.000"],
    ],
    widths=[1.0, 4.4, 6.6, 4.0],
)
para("Biaya berulang (recurring) untuk Opsi A:", bold=True)
table(
    ["Komponen", "Deskripsi", "Estimasi Biaya (IDR)"],
    [
        ["Maintenance & Support Tahunan",
         "Berlaku setelah masa garansi; opsional",
         "10.000.000 / tahun"],
        ["Server / VPS",
         "4 vCPU, 8 GB RAM (jika disediakan oleh kami)",
         "1.500.000 - 2.500.000 / bulan"],
    ],
    widths=[4.5, 7.5, 4.0],
)

h2("Opsi B — Langganan (Subscription)")
para("Model langganan sudah mencakup hosting, pemeliharaan, pembaruan, dan "
     "dukungan teknis — tanpa biaya awal besar.")
table(
    ["Paket", "Cakupan", "Estimasi Biaya (IDR)"],
    [
        ["Langganan Bulanan",
         "500 user, hosting, maintenance, update, support",
         "5.000.000 / bulan"],
        ["Langganan Tahunan",
         "500 user, hosting, maintenance, update, support (hemat ~13%)",
         "52.000.000 / tahun"],
    ],
    widths=[4.0, 8.0, 4.0],
)

h2("Termin Pembayaran (Opsi A)")
bullet("50% — uang muka saat penandatanganan kontrak (Down Payment);")
bullet("40% — setelah UAT (User Acceptance Test) disetujui;")
bullet("10% — saat go-live / serah-terima akhir.")

para("")
para("Perbandingan dengan paket lengkap:", bold=True)
table(
    ["Paket", "Cakupan", "Estimasi One-Time (IDR)"],
    [
        ["SDK Voice Call (proposal ini)", "Panggilan suara saja", "69.000.000"],
        ["SDK Voice Call + Video Meeting",
         "Panggilan suara + rapat video lengkap", "107.000.000"],
    ],
    widths=[5.5, 6.5, 4.0],
)

# ===========================================================================
# 8. MASA GARANSI
# ===========================================================================
h1("8. Masa Garansi (Warranty)")
para("Solusi diberikan garansi selama 6 (enam) bulan terhitung sejak tanggal "
     "serah-terima / go-live.", bold=True)

h2("8.1. Cakupan Garansi")
bullet("Perbaikan bug, error, dan kegagalan fungsi tanpa biaya tambahan;")
bullet("Perbaikan gangguan yang menyebabkan SDK / API tidak berjalan sesuai "
       "spesifikasi yang disepakati;")
bullet("Dukungan teknis melalui email / kanal komunikasi yang disepakati;")
bullet("Pendampingan terkait konfigurasi sistem yang telah diserahkan.")

h2("8.2. SLA Selama Masa Garansi")
table(
    ["Tingkat Prioritas", "Contoh Masalah", "Waktu Respon", "Target Penyelesaian"],
    [
        ["Kritis", "Layanan panggilan tidak dapat diakses total",
         "Maks. 4 jam kerja", "Maks. 1 hari kerja"],
        ["Tinggi", "Gagal memulai/menerima panggilan",
         "Maks. 8 jam kerja", "Maks. 2 hari kerja"],
        ["Sedang/Rendah", "Bug minor, masalah non-kritis",
         "Maks. 2 hari kerja", "Sesuai kesepakatan"],
    ],
    widths=[2.8, 5.2, 3.5, 4.5],
)

h2("8.3. Tidak Termasuk Garansi")
bullet("Penambahan fitur baru atau perubahan ruang lingkup di luar kesepakatan;")
bullet("Permintaan fitur Video Meeting / berbagi layar / perekaman "
       "(merupakan paket terpisah);")
bullet("Kerusakan akibat modifikasi kode oleh pihak di luar tim kami;")
bullet("Gangguan akibat infrastruktur di luar kendali kami (jaringan internet, "
       "listrik, kegagalan hardware, layanan cloud);")
bullet("Kesalahan penggunaan atau konfigurasi ulang tanpa koordinasi;")
bullet("Serangan keamanan akibat kelalaian pengelolaan kredensial oleh klien.")

h2("8.4. Setelah Masa Garansi")
para("Setelah 6 bulan, dukungan dan pemeliharaan dapat dilanjutkan melalui "
     "kontrak Maintenance & Support tahunan (lihat Bagian 7) yang mencakup "
     "pembaruan keamanan, pemantauan, serta dukungan teknis berkelanjutan.")

doc.add_page_break()

# ===========================================================================
# 9. LINGKUP & ASUMSI
# ===========================================================================
h1("9. Lingkup Pekerjaan & Asumsi")

h2("9.1. Termasuk dalam Penawaran")
for f in [
    "Penyediaan SDK Voice Call (@teleconf/voicecall-sdk)",
    "Layanan backend voice call + REST API v1",
    "Instalasi & konfigurasi pada 1 (satu) server produksi",
    "Dokumentasi OpenAPI / Swagger",
    "Pendampingan integrasi SDK ke 1 (satu) aplikasi pihak ketiga",
    "Pelatihan untuk tim pengembang",
    "Dokumentasi teknis & panduan penggunaan",
    "Garansi 6 bulan",
]:
    bullet(f)

h2("9.2. Tidak Termasuk")
for f in [
    "Fitur Video Meeting, berbagi layar, perekaman (paket terpisah)",
    "Pengadaan server / VPS / domain (kecuali disepakati pada Opsi terkait)",
    "Pengembangan antarmuka / aplikasi pihak ketiga itu sendiri",
    "Kustomisasi fitur di luar daftar pada Bagian 3",
    "Biaya langganan layanan cloud pihak ketiga",
]:
    bullet(f)

h2("9.3. Asumsi")
for f in [
    "Klien menyediakan akses server dengan IP publik dan domain",
    "Server memenuhi spesifikasi minimum pada Bagian 6",
    "Tersedia koneksi internet yang memadai dan stabil",
    "Klien menunjuk PIC teknis selama implementasi",
    "Aplikasi pihak ketiga berbasis web (Next.js / React) atau web view",
    "Sertifikat TLS menggunakan Let's Encrypt (gratis) atau disediakan klien",
]:
    bullet(f)

# ===========================================================================
# 10. TIMELINE
# ===========================================================================
h1("10. Estimasi Jadwal Implementasi")
para("Total estimasi waktu implementasi: 3-4 minggu kerja sejak kontrak "
     "ditandatangani dan prasyarat terpenuhi.")
table(
    ["Tahap", "Aktivitas", "Estimasi Durasi"],
    [
        ["1", "Persiapan, kick-off, penyiapan server", "Minggu 1"],
        ["2", "Instalasi, konfigurasi, deployment backend & API", "Minggu 1-2"],
        ["3", "Integrasi SDK ke aplikasi pihak ketiga", "Minggu 2-3"],
        ["4", "UAT (User Acceptance Test) & perbaikan", "Minggu 3"],
        ["5", "Pelatihan, go-live, serah-terima", "Minggu 4"],
    ],
    widths=[1.6, 9.4, 5.0],
)

# ===========================================================================
# 11. PENUTUP
# ===========================================================================
h1("11. Penutup")
para(
    "Demikian proposal ini kami sampaikan. SDK Voice Call dirancang sebagai "
    "komponen panggilan suara yang ringan, aman, dan cepat diintegrasikan ke "
    "dalam aplikasi Anda. Kami siap mendiskusikan penyesuaian ruang lingkup "
    "maupun harga sesuai kebutuhan, termasuk opsi peningkatan (upgrade) ke "
    "paket lengkap dengan Video Meeting di kemudian hari."
)
para("Atas perhatian dan kepercayaan Bapak/Ibu, kami ucapkan terima kasih.")

for _ in range(2):
    doc.add_paragraph()

table(
    ["Hormat kami,", "Menyetujui,"],
    [
        ["\n\n\n( ______________________ )\n[Nama Perusahaan Anda]",
         "\n\n\n( ______________________ )\n[Nama Klien / Instansi]"],
    ],
    widths=[8, 8],
)

doc.add_paragraph()
para("Catatan: seluruh angka biaya pada dokumen ini merupakan estimasi dan "
     "tidak mengikat sampai dituangkan dalam kontrak resmi.",
     italic=True, size=8.5, color=GREY)

out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   "Proposal-SDK-VoiceCall-Only.docx")
doc.save(out)
print("Dokumen tersimpan:", out)
