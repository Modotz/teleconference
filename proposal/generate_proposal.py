# -*- coding: utf-8 -*-
"""
Generator dokumen Word: Proposal Penjualan SDK Voice Call & Video Meeting.
Jalankan:  python generate_proposal.py
Output  :  Proposal-SDK-VoiceCall-VideoMeeting.docx
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0x25, 0x63, 0xEB)
GREY = RGBColor(0x55, 0x55, 0x55)
HEADER_FILL = "1F3A5F"
ALT_FILL = "EEF2F7"

doc = Document()

# ---- Page / base style ----------------------------------------------------
section = doc.sections[0]
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.2)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def h1(text):
    p = doc.add_heading(level=1)
    run = p.add_run(text)
    run.font.color.rgb = NAVY
    run.font.size = Pt(16)
    return p


def h2(text):
    p = doc.add_heading(level=2)
    run = p.add_run(text)
    run.font.color.rgb = ACCENT
    run.font.size = Pt(13)
    return p


def para(text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align is not None:
        p.alignment = align
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def mono(lines):
    p = doc.add_paragraph()
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        shade_cell(hdr[i], HEADER_FILL)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            run = cells[ci].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            if ri % 2 == 1:
                shade_cell(cells[ci], ALT_FILL)
    if widths:
        for row in t.rows:
            for ci, w in enumerate(widths):
                row.cells[ci].width = Cm(w)
    return t


def spacer(pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pt)


# ===========================================================================
# COVER
# ===========================================================================
for _ in range(3):
    doc.add_paragraph()

para("PROPOSAL PENAWARAN", bold=True, size=14, color=ACCENT,
     align=WD_ALIGN_PARAGRAPH.CENTER)
para("Solusi SDK Voice Call & Video Meeting", bold=True, size=26, color=NAVY,
     align=WD_ALIGN_PARAGRAPH.CENTER)
para("Platform Komunikasi Real-Time — Kapasitas 500 Pengguna", size=13,
     color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(3):
    doc.add_paragraph()

info = table(
    ["Keterangan", "Detail"],
    [
        ["Diajukan untuk", "[Nama Klien / Instansi]"],
        ["Disusun oleh", "[Nama Perusahaan Anda]"],
        ["Nomor Dokumen", "PROP/VC-VM/2026/001"],
        ["Tanggal", "Mei 2026"],
        ["Versi Dokumen", "1.0"],
        ["Masa Berlaku Penawaran", "30 hari kalender sejak tanggal terbit"],
    ],
    widths=[6, 10],
)

for _ in range(2):
    doc.add_paragraph()
para("Dokumen ini bersifat rahasia dan ditujukan khusus untuk pihak penerima. "
     "Dilarang menggandakan atau menyebarluaskan tanpa izin tertulis.",
     italic=True, size=9, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ===========================================================================
# 1. RINGKASAN EKSEKUTIF
# ===========================================================================
h1("1. Ringkasan Eksekutif")
para(
    "Kami menawarkan solusi komunikasi real-time terintegrasi yang terdiri dari "
    "SDK Voice Call dan aplikasi Video Meeting. Solusi ini dirancang untuk "
    "melayani hingga 500 pengguna dan dapat di-deploy secara mandiri "
    "(self-hosted) di server milik klien maupun cloud."
)
para(
    "Keunggulan utama solusi ini adalah SDK (Software Development Kit) yang "
    "memungkinkan aplikasi pihak ketiga — termasuk aplikasi berbasis Next.js / "
    "React — menanamkan fitur panggilan suara dan video langsung ke dalam "
    "produk mereka, tanpa pengguna perlu membuat akun terpisah. Seluruh media "
    "(audio/video) berjalan di atas standar WebRTC dengan arsitektur SFU "
    "(Selective Forwarding Unit) sehingga efisien dan skalabel."
)
bullet("aplikasi Video Meeting siap pakai dengan fitur lengkap;", "Yang ditawarkan: ")
bullet("SDK Voice Call yang dapat diintegrasikan ke aplikasi pihak ketiga;")
bullet("REST API publik beserta dokumentasi OpenAPI/Swagger;")
bullet("dukungan instalasi, integrasi, pelatihan, dan garansi.")

# ===========================================================================
# 2. LATAR BELAKANG & TUJUAN
# ===========================================================================
h1("2. Latar Belakang & Tujuan")
para(
    "Kebutuhan komunikasi jarak jauh — rapat daring, panggilan suara, dan "
    "kolaborasi tim — terus meningkat. Banyak organisasi membutuhkan solusi "
    "yang (1) dapat dikendalikan penuh secara mandiri, (2) menjaga privasi "
    "data karena tidak bergantung pada layanan pihak ketiga, dan (3) dapat "
    "ditanamkan ke dalam aplikasi internal yang sudah ada."
)
para("Tujuan penyediaan solusi ini:", bold=True)
bullet("menyediakan platform Video Meeting mandiri untuk hingga 500 pengguna;")
bullet("menyediakan SDK Voice Call yang mudah diintegrasikan oleh tim "
       "pengembang aplikasi pihak ketiga;")
bullet("memastikan keamanan komunikasi melalui enkripsi standar WebRTC;")
bullet("memberikan kepemilikan penuh atas data dan infrastruktur kepada klien.")

# ===========================================================================
# 3. FITUR APLIKASI
# ===========================================================================
h1("3. Fitur Aplikasi")

h2("3.1. Video Meeting")
for f in [
    "Konferensi video & audio multi-peserta dengan arsitektur SFU",
    "Berbagi layar (screen sharing) lengkap dengan audio sistem",
    "Perekaman rapat (recording) sisi klien",
    "Chat dalam ruang: teks, emoji, gambar, dan file (hingga 25 MB)",
    "Pengaturan tata letak: Grid, Spotlight, dan Sidebar",
    "Sematkan peserta (pin) ke tampilan utama",
    "Angkat tangan (raise hand)",
    "Deteksi pembicara aktif (active speaker detection)",
    "Virtual background: efek blur atau ganti latar dengan gambar",
    "Penjadwalan rapat (schedule meeting) dan tautan undangan (share link)",
    "Antarmuka responsif untuk desktop maupun perangkat mobile",
]:
    bullet(f)

h2("3.2. Voice Call")
for f in [
    "Panggilan suara 1-lawan-1 dan panggilan grup",
    "Panggilan di dalam aplikasi maupun melalui SDK pihak ketiga",
    "Kontrol mute / unmute dengan indikator pembicara aktif",
    "Notifikasi panggilan masuk (incoming call)",
    "Audio jernih dengan codec Opus",
]:
    bullet(f)

h2("3.3. Chat / Pesan")
for f in [
    "Chat pribadi dan grup bergaya populer (mirip WhatsApp)",
    "Daftar kontak diambil dari basis data pengguna",
    "Emoji, kirim gambar & file, serta pesan suara (voice message)",
    "Indikator sedang mengetik (typing indicator)",
    "Tanda pesan telah dibaca (read receipts)",
    "Tambah / keluarkan anggota grup",
    "Riwayat percakapan tersimpan permanen",
]:
    bullet(f)

h2("3.4. SDK & Integrasi Pihak Ketiga")
for f in [
    "SDK JavaScript / TypeScript: paket @teleconf/voicecall-sdk",
    "Komponen siap pakai untuk Next.js / React serta React Hook (headless)",
    "REST API publik (v1) dengan autentikasi API Key",
    "Access token untuk peserta tamu — tanpa perlu membuat akun",
    "Dokumentasi interaktif OpenAPI 3.0 / Swagger UI",
    "Mudah diintegrasikan: cukup mint token di sisi server, embed di klien",
]:
    bullet(f)

h2("3.5. Keamanan & Autentikasi")
for f in [
    "Login berbasis JWT dengan kata sandi terenkripsi (bcrypt)",
    "Media audio/video terenkripsi (DTLS-SRTP, standar WebRTC)",
    "Seluruh komunikasi melalui HTTPS / WSS",
    "API Key disimpan dalam bentuk hash; access token berumur pendek",
]:
    bullet(f)

doc.add_page_break()

# ===========================================================================
# 4. ARSITEKTUR & TOPOLOGI
# ===========================================================================
h1("4. Arsitektur & Topologi Sistem")
para(
    "Sistem menggunakan arsitektur SFU (Selective Forwarding Unit). Setiap "
    "peserta hanya mengirim satu aliran media ke server, lalu server "
    "meneruskan (forward) aliran tersebut ke peserta lain. Pendekatan ini jauh "
    "lebih efisien dan skalabel dibanding arsitektur mesh (peer-to-peer "
    "penuh), terutama untuk rapat dengan banyak peserta."
)

h2("4.1. Diagram Topologi")
mono([
    "        +----------------------- KLIEN -------------------------+",
    "        |                                                       |",
    "   [ Browser    ]   [ Browser   ]   [ Aplikasi Pihak Ketiga ]   ",
    "   [ Desktop    ]   [ Mobile/HP ]   [ Next.js + Voice SDK   ]   ",
    "        |                |                     |               ",
    "        +--------+-------+----------+----------+               ",
    "                 |  HTTPS / WSS / WebRTC (UDP+TCP)              ",
    "                 v                                             ",
    "        =================  INTERNET  =======================   ",
    "                 |                                             ",
    "                 v                                             ",
    "   +---------------------------------------------------------+ ",
    "   |               SERVER (Cloud / On-Premise)               | ",
    "   |                                                         | ",
    "   |   [ Nginx ]  reverse proxy + TLS (HTTPS/WSS)             | ",
    "   |       |                                                 | ",
    "   |       +---> [ Frontend  ]  Next.js (Web UI)              | ",
    "   |       +---> [ Backend   ]  Node.js + Express + Socket.io | ",
    "   |       |     [ Mediasoup ]  WebRTC SFU (media server)     | ",
    "   |       +---> [ REST API v1 ] untuk SDK pihak ketiga       | ",
    "   |                     |                                   | ",
    "   |                     v                                   | ",
    "   |             [ PostgreSQL ]  basis data                  | ",
    "   +---------------------------------------------------------+ ",
])
para("Catatan: media audio/video mengalir langsung antara klien dan Mediasoup "
     "(SFU) melalui WebRTC; jalur signaling memakai WebSocket (Socket.io).",
     italic=True, size=9, color=GREY)

h2("4.2. Komponen Sistem")
table(
    ["Lapisan", "Komponen", "Fungsi"],
    [
        ["Klien", "Browser & Aplikasi Pihak Ketiga",
         "Antarmuka pengguna; menjalankan SDK voice/video"],
        ["Gateway", "Nginx", "Reverse proxy, terminasi TLS, load balancing"],
        ["Aplikasi", "Frontend Next.js", "Antarmuka Video Meeting & Chat"],
        ["Aplikasi", "Backend Node.js", "Signaling, autentikasi, REST API"],
        ["Media", "Mediasoup (SFU)", "Penerusan aliran audio/video WebRTC"],
        ["Data", "PostgreSQL", "Pengguna, ruang, pesan, jadwal, API client"],
    ],
    widths=[2.6, 5.2, 8.2],
)

h2("4.3. Aliran Integrasi SDK Pihak Ketiga")
mono([
    "  Backend Partner  --X-Api-Key-->  POST /api/v1/calls          -> callId",
    "  Backend Partner  --X-Api-Key-->  POST /api/v1/calls/:id/token -> token",
    "  Browser Partner  --accessToken-> SDK Voice Call --WebRTC--> Mediasoup",
])
para("API Key hanya digunakan di sisi server partner. Browser hanya menerima "
     "access token berumur pendek yang terbatas pada satu panggilan.",
     italic=True, size=9, color=GREY)

doc.add_page_break()

# ===========================================================================
# 5. TEKNOLOGI YANG DIGUNAKAN
# ===========================================================================
h1("5. Teknologi yang Digunakan")
para("Solusi dibangun dengan teknologi modern, open standard, dan teruji di "
     "lingkungan produksi.")
table(
    ["Kategori", "Teknologi", "Keterangan"],
    [
        ["Frontend", "Next.js 15, React 19, TypeScript, TailwindCSS",
         "Antarmuka web responsif"],
        ["Backend", "Node.js 20+, Express, Socket.io",
         "API, signaling real-time"],
        ["Media Server", "Mediasoup (WebRTC SFU)",
         "Penerusan media efisien & skalabel"],
        ["Basis Data", "PostgreSQL 16 + Prisma ORM",
         "Penyimpanan data relasional"],
        ["Real-Time", "WebRTC, WebSocket",
         "Audio/video & signaling latensi rendah"],
        ["Virtual Background", "Google MediaPipe",
         "Segmentasi latar secara lokal di browser"],
        ["SDK", "TypeScript, mediasoup-client, tsup",
         "Paket terdistribusi untuk pihak ketiga"],
        ["Dokumentasi API", "OpenAPI 3.0, Swagger UI",
         "Referensi REST API interaktif"],
        ["Deployment", "Docker, Nginx, Let's Encrypt",
         "Kontainerisasi, reverse proxy, sertifikat TLS"],
        ["Keamanan", "JWT, bcrypt, DTLS-SRTP",
         "Autentikasi & enkripsi end-to-end media"],
    ],
    widths=[3.0, 6.0, 7.0],
)

# ===========================================================================
# 6. SPESIFIKASI SERVER
# ===========================================================================
h1("6. Spesifikasi Server")
para(
    "Spesifikasi berikut diperhitungkan untuk 500 pengguna terdaftar dengan "
    "asumsi puncak penggunaan bersamaan (concurrent) sekitar 100-150 pengguna "
    "aktif dalam panggilan/rapat. Beban utama server adalah CPU (penerusan "
    "media) dan bandwidth jaringan."
)

h2("6.1. Server Produksi (Direkomendasikan)")
table(
    ["Komponen", "Spesifikasi Minimum", "Spesifikasi Direkomendasikan"],
    [
        ["CPU", "6 vCPU", "8 vCPU (Intel Xeon / AMD EPYC generasi baru)"],
        ["RAM", "12 GB", "16 GB"],
        ["Penyimpanan", "120 GB SSD", "160 GB SSD NVMe"],
        ["Jaringan", "1 Gbps", "1 Gbps, kuota bandwidth tinggi / unmetered"],
        ["Sistem Operasi", "Ubuntu Server 22.04 LTS", "Ubuntu Server 22.04 LTS"],
        ["IP Publik", "1 IP statis", "1 IP statis"],
        ["Port", "80, 443, 40000-40100 (UDP+TCP)",
         "80, 443, 40000-40100 (UDP+TCP)"],
    ],
    widths=[3.2, 5.4, 7.4],
)

h2("6.2. Catatan Kapasitas & Skalabilitas")
bullet("Bandwidth adalah faktor biaya utama. Rapat video 8 peserta dapat "
       "menghasilkan puluhan Mbps lalu lintas server; pastikan kuota memadai "
       "atau gunakan layanan unmetered.")
bullet("Mediasoup berjalan satu worker per inti CPU. Penambahan inti CPU "
       "meningkatkan kapasitas peserta secara linear.")
bullet("Untuk pertumbuhan di atas 500 pengguna, sistem dapat di-scale "
       "horizontal dengan menambah server media (multi-node Mediasoup).")
bullet("Panggilan suara (voice-only) jauh lebih ringan dari video; kapasitas "
       "voice call bisa beberapa kali lipat lebih besar pada spesifikasi sama.")
bullet("Database PostgreSQL dapat ditempatkan pada server yang sama; untuk "
       "keandalan lebih tinggi disarankan database terkelola terpisah.")

h2("6.3. Lingkungan Pengujian (Opsional)")
table(
    ["Komponen", "Spesifikasi"],
    [
        ["CPU", "2-4 vCPU"],
        ["RAM", "4-8 GB"],
        ["Penyimpanan", "40 GB SSD"],
        ["Kegunaan", "Staging / UAT / pelatihan"],
    ],
    widths=[5.0, 11.0],
)

doc.add_page_break()

# ===========================================================================
# 7. ESTIMASI BIAYA
# ===========================================================================
h1("7. Estimasi Biaya")
para(
    "Tersedia dua skema kepemilikan. Klien dapat memilih sesuai kebutuhan "
    "anggaran dan model operasional.", bold=True)
para("Seluruh harga dalam Rupiah (IDR), belum termasuk PPN. Angka bersifat "
     "estimasi dan dapat disesuaikan berdasarkan kesepakatan akhir serta "
     "ruang lingkup final.", italic=True, size=9, color=GREY)

h2("Opsi A — Beli Putus (Perpetual License)")
table(
    ["No", "Komponen", "Deskripsi", "Estimasi Biaya (IDR)"],
    [
        ["1", "Lisensi Software",
         "SDK Voice Call & Video Meeting, kapasitas 500 user, lisensi perpetual",
         "75.000.000"],
        ["2", "Implementasi & Deployment",
         "Instalasi, konfigurasi server, hardening, go-live",
         "15.000.000"],
        ["3", "Integrasi SDK",
         "Pendampingan integrasi SDK ke aplikasi pihak ketiga",
         "12.000.000"],
        ["4", "Pelatihan & Dokumentasi",
         "Pelatihan admin & developer, serah-terima dokumen",
         "5.000.000"],
        ["", "TOTAL BIAYA AWAL (One-Time)", "", "107.000.000"],
    ],
    widths=[1.0, 4.2, 6.8, 4.0],
)
para("Biaya berulang (recurring) untuk Opsi A:", bold=True)
table(
    ["Komponen", "Deskripsi", "Estimasi Biaya (IDR)"],
    [
        ["Maintenance & Support Tahunan",
         "Berlaku setelah masa garansi; opsional",
         "18.000.000 / tahun"],
        ["Server / VPS",
         "8 vCPU, 16 GB RAM (jika disediakan oleh kami)",
         "3.000.000 - 4.000.000 / bulan"],
    ],
    widths=[4.5, 7.5, 4.0],
)

h2("Opsi B — Langganan (Subscription)")
para("Model langganan sudah mencakup hosting, pemeliharaan, pembaruan, dan "
     "dukungan teknis — tanpa biaya awal besar.")
table(
    ["Paket", "Cakupan", "Estimasi Biaya (IDR)"],
    [
        ["Langganan Bulanan",
         "500 user, hosting, maintenance, update, support",
         "9.000.000 / bulan"],
        ["Langganan Tahunan",
         "500 user, hosting, maintenance, update, support (hemat ~12%)",
         "95.000.000 / tahun"],
    ],
    widths=[4.0, 8.0, 4.0],
)

h2("Termin Pembayaran (Opsi A)")
bullet("50% — uang muka saat penandatanganan kontrak (Down Payment);", "")
bullet("40% — setelah UAT (User Acceptance Test) disetujui;")
bullet("10% — saat go-live / serah-terima akhir.")

# ===========================================================================
# 8. MASA GARANSI
# ===========================================================================
h1("8. Masa Garansi (Warranty)")
para("Solusi diberikan garansi selama 6 (enam) bulan terhitung sejak tanggal "
     "serah-terima / go-live.", bold=True)

h2("8.1. Cakupan Garansi")
bullet("Perbaikan bug, error, dan kegagalan fungsi tanpa biaya tambahan;")
bullet("Perbaikan gangguan yang menyebabkan fitur tidak berjalan sesuai "
       "spesifikasi yang disepakati;")
bullet("Dukungan teknis melalui email / kanal komunikasi yang disepakati;")
bullet("Pendampingan terkait konfigurasi sistem yang telah diserahkan.")

h2("8.2. SLA Selama Masa Garansi")
table(
    ["Tingkat Prioritas", "Contoh Masalah", "Waktu Respon", "Target Penyelesaian"],
    [
        ["Kritis", "Layanan tidak dapat diakses total",
         "Maks. 4 jam kerja", "Maks. 1 hari kerja"],
        ["Tinggi", "Fitur utama gagal (mis. tidak bisa join call)",
         "Maks. 8 jam kerja", "Maks. 2 hari kerja"],
        ["Sedang/Rendah", "Bug minor, masalah tampilan",
         "Maks. 2 hari kerja", "Sesuai kesepakatan"],
    ],
    widths=[2.8, 5.2, 3.5, 4.5],
)

h2("8.3. Tidak Termasuk Garansi")
bullet("Penambahan fitur baru atau perubahan ruang lingkup di luar kesepakatan;")
bullet("Kerusakan akibat modifikasi kode oleh pihak di luar tim kami;")
bullet("Gangguan akibat infrastruktur di luar kendali kami (jaringan internet, "
       "listrik, kegagalan hardware pihak ketiga, layanan cloud);")
bullet("Kesalahan penggunaan atau konfigurasi ulang tanpa koordinasi;")
bullet("Serangan keamanan akibat kelalaian pengelolaan kredensial oleh klien.")

h2("8.4. Setelah Masa Garansi")
para("Setelah 6 bulan, dukungan dan pemeliharaan dapat dilanjutkan melalui "
     "kontrak Maintenance & Support tahunan (lihat Bagian 7) yang mencakup "
     "pembaruan keamanan, pemantauan, serta dukungan teknis berkelanjutan.")

doc.add_page_break()

# ===========================================================================
# 9. LINGKUP & ASUMSI
# ===========================================================================
h1("9. Lingkup Pekerjaan & Asumsi")

h2("9.1. Termasuk dalam Penawaran")
for f in [
    "Penyediaan aplikasi Video Meeting dan SDK Voice Call",
    "Instalasi & konfigurasi pada 1 (satu) server produksi",
    "Integrasi REST API & dokumentasi OpenAPI",
    "Pendampingan integrasi SDK ke 1 (satu) aplikasi pihak ketiga",
    "Pelatihan untuk administrator dan tim pengembang",
    "Dokumentasi teknis & panduan penggunaan",
    "Garansi 6 bulan",
]:
    bullet(f)

h2("9.2. Tidak Termasuk")
for f in [
    "Pengadaan server / VPS / domain (kecuali disepakati pada Opsi terkait)",
    "Pengembangan aplikasi pihak ketiga itu sendiri",
    "Kustomisasi fitur di luar daftar pada Bagian 3",
    "Migrasi data dari sistem lama",
    "Biaya langganan layanan cloud pihak ketiga",
]:
    bullet(f)

h2("9.3. Asumsi")
for f in [
    "Klien menyediakan akses server dengan IP publik dan domain",
    "Server memenuhi spesifikasi minimum pada Bagian 6",
    "Tersedia koneksi internet yang memadai dan stabil",
    "Klien menunjuk PIC teknis selama implementasi",
    "Sertifikat TLS menggunakan Let's Encrypt (gratis) atau disediakan klien",
]:
    bullet(f)

# ===========================================================================
# 10. TIMELINE
# ===========================================================================
h1("10. Estimasi Jadwal Implementasi")
para("Total estimasi waktu implementasi: 4-6 minggu kerja sejak kontrak "
     "ditandatangani dan prasyarat terpenuhi.")
table(
    ["Tahap", "Aktivitas", "Estimasi Durasi"],
    [
        ["1", "Persiapan, kick-off, penyiapan server", "Minggu 1"],
        ["2", "Instalasi, konfigurasi, deployment", "Minggu 2"],
        ["3", "Integrasi SDK & penyesuaian", "Minggu 3-4"],
        ["4", "UAT (User Acceptance Test) & perbaikan", "Minggu 5"],
        ["5", "Pelatihan, go-live, serah-terima", "Minggu 6"],
    ],
    widths=[1.6, 9.4, 5.0],
)

# ===========================================================================
# 11. PENUTUP
# ===========================================================================
h1("11. Penutup")
para(
    "Demikian proposal ini kami sampaikan. Solusi SDK Voice Call & Video "
    "Meeting ini dirancang untuk memberikan kendali penuh, keamanan, dan "
    "fleksibilitas integrasi bagi organisasi Anda. Kami siap mendiskusikan "
    "penyesuaian ruang lingkup maupun harga sesuai kebutuhan."
)
para("Atas perhatian dan kepercayaan Bapak/Ibu, kami ucapkan terima kasih.")

for _ in range(2):
    doc.add_paragraph()

sign = table(
    ["Hormat kami,", "Menyetujui,"],
    [
        ["\n\n\n( ______________________ )\n[Nama Perusahaan Anda]",
         "\n\n\n( ______________________ )\n[Nama Klien / Instansi]"],
    ],
    widths=[8, 8],
)

doc.add_paragraph()
para("Catatan: seluruh angka biaya pada dokumen ini merupakan estimasi dan "
     "tidak mengikat sampai dituangkan dalam kontrak resmi.",
     italic=True, size=8.5, color=GREY)

# ---- Save -----------------------------------------------------------------
out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   "Proposal-SDK-VoiceCall-VideoMeeting.docx")
doc.save(out)
print("Dokumen tersimpan:", out)
