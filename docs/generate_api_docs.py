# -*- coding: utf-8 -*-
"""
Generator: API Documentation — SDK Voice Call.
Jalankan:  python generate_api_docs.py
Output  :  API-Documentation-SDK-VoiceCall.docx
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0x25, 0x63, 0xEB)
GREEN = RGBColor(0x15, 0x80, 0x3D)
GREY = RGBColor(0x55, 0x55, 0x55)
HEADER_FILL = "1F3A5F"
ALT_FILL = "EEF2F7"
CODE_FILL = "F2F3F5"

doc = Document()
s = doc.sections[0]
s.left_margin = s.right_margin = Cm(2.4)
s.top_margin = s.bottom_margin = Cm(2.2)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def h1(text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.color.rgb = NAVY
    r.font.size = Pt(16)
    return p


def h2(text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    r.font.color.rgb = ACCENT
    r.font.size = Pt(13)
    return p


def h3(text):
    p = doc.add_heading(level=3)
    r = p.add_run(text)
    r.font.color.rgb = NAVY
    r.font.size = Pt(11.5)
    return p


def para(text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if align is not None:
        p.alignment = align
    return p


def bullet(text, prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if prefix:
        p.add_run(prefix).bold = True
    p.add_run(text)
    return p


def code(text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), CODE_FILL)
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    return p


def note(text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFF7E0")
    pPr.append(shd)
    r = p.add_run("  Catatan: " + text)
    r.font.size = Pt(9.5)
    r.italic = True
    return p


def endpoint(method, path):
    """A coloured 'METHOD /path' line."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "E8EEF7")
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    rm = p.add_run(" " + method + " ")
    rm.bold = True
    rm.font.name = "Consolas"
    rm.font.size = Pt(10)
    rm.font.color.rgb = GREEN if method == "GET" else ACCENT
    rp = p.add_run(path + " ")
    rp.font.name = "Consolas"
    rp.font.size = Pt(10)
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(htext)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9.5)
        shade_cell(c, HEADER_FILL)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            r = cells[ci].paragraphs[0].add_run(str(val))
            r.font.size = Pt(9)
            if ri % 2 == 1:
                shade_cell(cells[ci], ALT_FILL)
    if widths:
        for row in t.rows:
            for ci, w in enumerate(widths):
                row.cells[ci].width = Cm(w)
    return t


# ===========================================================================
# COVER
# ===========================================================================
for _ in range(4):
    doc.add_paragraph()
para("API DOCUMENTATION", bold=True, size=15, color=ACCENT,
     align=WD_ALIGN_PARAGRAPH.CENTER)
para("SDK Voice Call", bold=True, size=30, color=NAVY,
     align=WD_ALIGN_PARAGRAPH.CENTER)
para("Referensi REST API & SDK JavaScript / TypeScript", size=13, color=GREY,
     align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(3):
    doc.add_paragraph()
table(
    ["Keterangan", "Detail"],
    [
        ["Produk", "SDK Voice Call (@teleconf/voicecall-sdk)"],
        ["Versi API", "v1"],
        ["Versi SDK", "0.1.0"],
        ["Versi Dokumen", "1.0"],
        ["Tanggal", "Mei 2026"],
        ["Dokumentasi Interaktif", "https://[server]/api/v1/docs (Swagger UI)"],
    ],
    widths=[5, 11],
)
doc.add_page_break()

# ===========================================================================
# DAFTAR ISI
# ===========================================================================
h1("Daftar Isi")
for item in [
    "1. Pendahuluan",
    "2. Autentikasi",
    "3. Konvensi & Penanganan Error",
    "4. REST API v1",
    "   4.1. POST /api/v1/calls",
    "   4.2. GET /api/v1/calls/{id}",
    "   4.3. POST /api/v1/calls/{id}/token",
    "   4.4. POST /api/v1/calls/{id}/end",
    "5. SDK JavaScript / TypeScript",
    "   5.1. VoiceCallClient",
    "   5.2. Hook useVoiceCall",
    "   5.3. Komponen VoiceCallRoom",
    "6. Definisi Tipe Data",
    "7. Contoh Integrasi End-to-End",
    "8. Daftar Kode Status",
]:
    doc.add_paragraph().add_run(item).font.size = Pt(11)
doc.add_page_break()

# ===========================================================================
# 1. PENDAHULUAN
# ===========================================================================
h1("1. Pendahuluan")
para("Dokumen ini adalah referensi teknis lengkap untuk SDK Voice Call, "
     "terdiri atas dua bagian:")
bullet("REST API v1 — dipanggil dari server aplikasi Anda untuk membuat "
       "panggilan dan menerbitkan access token.", "REST API. ")
bullet("SDK JavaScript/TypeScript — dipakai di sisi browser untuk bergabung "
       "ke panggilan.", "SDK. ")
h2("1.1. Base URL")
para("Seluruh endpoint REST berada di bawah prefiks berikut:")
code("https://<host-server>/api/v1")
h2("1.2. Dokumentasi Interaktif")
para("Spesifikasi OpenAPI 3.0 dan antarmuka Swagger UI tersedia secara "
     "langsung pada server:")
bullet("Swagger UI : https://<host-server>/api/v1/docs")
bullet("Spesifikasi : https://<host-server>/api/v1/openapi.json")

# ===========================================================================
# 2. AUTENTIKASI
# ===========================================================================
h1("2. Autentikasi")
para("Terdapat dua mekanisme autentikasi yang berbeda peran.")

h2("2.1. API Key (REST API)")
para("Setiap permintaan ke REST API v1 wajib menyertakan header X-Api-Key. "
     "API Key bersifat rahasia dan HANYA boleh dipakai dari server aplikasi "
     "Anda.")
code("X-Api-Key: vc_3f9a1b2c4d5e6f7a8b9c0d1e2f3a4b5c6d7e8f9a")
note("API Key tidak boleh dikirim atau disimpan di sisi browser. Kebocoran "
     "API Key memungkinkan pihak lain membuat panggilan atas nama Anda.")

h2("2.2. Access Token (SDK)")
para("Access token adalah JWT berumur pendek yang diterbitkan oleh endpoint "
     "POST /calls/{id}/token. Token ini diberikan ke browser dan dipakai SDK "
     "untuk bergabung ke panggilan. Token memuat: callId, participantId, dan "
     "displayName. Masa berlaku standar 15 menit (jendela untuk join).")

# ===========================================================================
# 3. KONVENSI & ERROR
# ===========================================================================
h1("3. Konvensi & Penanganan Error")
bullet("Semua request & response berformat JSON (application/json).")
bullet("Body request POST dikirim sebagai JSON.")
bullet("Kode 2xx menandakan sukses; 4xx kesalahan klien; 5xx kesalahan server.")
para("Format response error selalu konsisten:")
code('{\n  "error": "Penjelasan singkat kesalahan"\n}')
para("Kode status yang umum:")
table(
    ["Kode", "Arti", "Penyebab Umum"],
    [
        ["200", "OK", "Permintaan berhasil"],
        ["201", "Created", "Sumber daya berhasil dibuat"],
        ["400", "Bad Request", "Parameter wajib tidak lengkap"],
        ["401", "Unauthorized", "API Key tidak ada / tidak valid"],
        ["404", "Not Found", "Call tidak ditemukan"],
        ["409", "Conflict", "Call sudah diakhiri"],
        ["500", "Server Error", "Kesalahan internal server"],
    ],
    widths=[2.0, 4.0, 10.0],
)

doc.add_page_break()

# ===========================================================================
# 4. REST API v1
# ===========================================================================
h1("4. REST API v1")
para("Seluruh endpoint berikut memerlukan header X-Api-Key.")

# --- 4.1 -------------------------------------------------------------------
h2("4.1. Membuat Call")
endpoint("POST", "/api/v1/calls")
para("Membuat sesi panggilan baru. Ruang media dibuat otomatis saat peserta "
     "pertama bergabung.")
para("Header", bold=True, size=10)
table(["Header", "Wajib", "Nilai"],
      [["X-Api-Key", "Ya", "API Key aplikasi Anda"]],
      widths=[4, 2.5, 9.5])
para("Request Body", bold=True, size=10)
para("Tidak ada.", italic=True, size=10)
para("Response 201 — Created", bold=True, size=10)
code('{\n'
     '  "callId": "8c1e4d2a-...-6caa620d377f",\n'
     '  "createdAt": "2026-05-22T08:30:00.000Z"\n'
     '}')
para("Contoh", bold=True, size=10)
code('curl -X POST https://voice.contoh.com/api/v1/calls \\\n'
     '  -H "X-Api-Key: vc_3f9a..."')

# --- 4.2 -------------------------------------------------------------------
h2("4.2. Status Call")
endpoint("GET", "/api/v1/calls/{id}")
para("Mengambil status panggilan beserta daftar peserta yang sedang "
     "terhubung.")
para("Path Parameter", bold=True, size=10)
table(["Parameter", "Tipe", "Keterangan"],
      [["id", "string (uuid)", "ID panggilan"]],
      widths=[3.5, 4.0, 8.5])
para("Response 200 — OK", bold=True, size=10)
code('{\n'
     '  "callId": "8c1e4d2a-...",\n'
     '  "createdAt": "2026-05-22T08:30:00.000Z",\n'
     '  "endedAt": null,\n'
     '  "active": true,\n'
     '  "participantCount": 2,\n'
     '  "participants": [\n'
     '    { "participantId": "p1", "displayName": "Budi" },\n'
     '    { "participantId": "p2", "displayName": "Sari" }\n'
     '  ]\n'
     '}')
para("Error", bold=True, size=10)
bullet("404 — Call tidak ditemukan / bukan milik aplikasi Anda.")

# --- 4.3 -------------------------------------------------------------------
h2("4.3. Menerbitkan Access Token")
endpoint("POST", "/api/v1/calls/{id}/token")
para("Menerbitkan satu access token untuk satu peserta tamu. Untuk panggilan "
     "grup, panggil endpoint ini sekali untuk tiap peserta dengan id yang "
     "sama.")
para("Path Parameter", bold=True, size=10)
table(["Parameter", "Tipe", "Keterangan"],
      [["id", "string (uuid)", "ID panggilan"]],
      widths=[3.5, 4.0, 8.5])
para("Request Body", bold=True, size=10)
table(
    ["Field", "Tipe", "Wajib", "Keterangan"],
    [
        ["displayName", "string", "Ya",
         "Nama peserta yang tampil ke peserta lain (maks. 80 karakter)"],
        ["externalId", "string", "Tidak",
         "ID stabil dari sistem Anda. Bila kosong, dibuat acak."],
    ],
    widths=[3.2, 2.4, 1.8, 8.6],
)
para("Response 201 — Created", bold=True, size=10)
code('{\n'
     '  "accessToken": "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9...",\n'
     '  "callId": "8c1e4d2a-...",\n'
     '  "participantId": "b9c0d1e2-...",\n'
     '  "expiresIn": 900\n'
     '}')
para("Error", bold=True, size=10)
bullet("400 — displayName tidak disertakan.")
bullet("404 — Call tidak ditemukan.")
bullet("409 — Call sudah diakhiri.")
para("Contoh", bold=True, size=10)
code('curl -X POST https://voice.contoh.com/api/v1/calls/8c1e.../token \\\n'
     '  -H "X-Api-Key: vc_3f9a..." \\\n'
     '  -H "Content-Type: application/json" \\\n'
     '  -d \'{"displayName":"Budi"}\'')

# --- 4.4 -------------------------------------------------------------------
h2("4.4. Mengakhiri Call")
endpoint("POST", "/api/v1/calls/{id}/end")
para("Menandai panggilan telah berakhir. Setelah ini tidak ada peserta baru "
     "yang dapat menerima token / bergabung.")
para("Response 200 — OK", bold=True, size=10)
code('{\n  "ok": true,\n  "callId": "8c1e4d2a-..."\n}')
para("Error", bold=True, size=10)
bullet("404 — Call tidak ditemukan.")

doc.add_page_break()

# ===========================================================================
# 5. SDK JS/TS
# ===========================================================================
h1("5. SDK JavaScript / TypeScript")
para("Paket: @teleconf/voicecall-sdk. Tiga antarmuka tersedia: kelas "
     "VoiceCallClient, hook useVoiceCall, dan komponen VoiceCallRoom.")
code("import {\n"
     "  VoiceCallClient,\n"
     "  useVoiceCall,\n"
     "  VoiceCallRoom,\n"
     "} from '@teleconf/voicecall-sdk';")

# --- 5.1 VoiceCallClient ---------------------------------------------------
h2("5.1. VoiceCallClient")
para("Kelas inti untuk mengendalikan panggilan tanpa React.")

h3("Constructor")
code("new VoiceCallClient({ serverUrl: string })")
table(["Opsi", "Tipe", "Keterangan"],
      [["serverUrl", "string", "Base URL server Voice Call"]],
      widths=[3.5, 3.0, 9.5])

h3("Properti")
table(["Properti", "Tipe", "Keterangan"],
      [["status", "VoiceCallStatus", "Status panggilan saat ini (read-only)"]],
      widths=[3.5, 4.5, 8.0])

h3("Method")
table(
    ["Method", "Mengembalikan", "Keterangan"],
    [
        ["join(accessToken: string)", "Promise<void>",
         "Minta izin mikrofon, sambungkan, dan gabung ke panggilan"],
        ["setMicEnabled(enabled: boolean)", "void",
         "Aktifkan / bisukan mikrofon lokal"],
        ["leave()", "void",
         "Keluar dari panggilan dan lepaskan semua sumber daya"],
        ["on(event, handler)", "() => void",
         "Berlangganan event; mengembalikan fungsi untuk berhenti berlangganan"],
    ],
    widths=[5.2, 3.4, 7.4],
)

h3("Event")
para("Berlangganan melalui client.on(namaEvent, handler).")
table(
    ["Event", "Argumen Handler", "Dipicu Saat"],
    [
        ["status", "VoiceCallStatus", "Status panggilan berubah"],
        ["participantJoined", "{ peerId, displayName }",
         "Peserta lain bergabung"],
        ["participantLeft", "peerId: string", "Peserta lain keluar"],
        ["remoteStream", "VoiceCallParticipant",
         "Audio peserta lain siap diputar"],
        ["activeSpeaker", "peerId | null",
         "Pembicara paling keras berubah"],
        ["error", "Error", "Terjadi kesalahan"],
    ],
    widths=[3.4, 5.0, 7.6],
)
para("Contoh", bold=True, size=10)
code(
    "const client = new VoiceCallClient({\n"
    "  serverUrl: 'https://voice.contoh.com',\n"
    "});\n\n"
    "client.on('status', (s) => console.log('status:', s));\n"
    "client.on('remoteStream', ({ displayName, stream }) => {\n"
    "  const audio = new Audio();\n"
    "  audio.srcObject = stream;\n"
    "  audio.play();\n"
    "});\n\n"
    "await client.join(accessToken);\n"
    "client.setMicEnabled(false);   // bisukan\n"
    "client.leave();"
)

# --- 5.2 useVoiceCall ------------------------------------------------------
h2("5.2. Hook useVoiceCall")
para("Hook React yang membungkus VoiceCallClient dan memutar audio peserta "
     "lain secara otomatis.")
code("const call = useVoiceCall({ serverUrl: string });")
h3("Nilai Kembalian")
table(
    ["Field", "Tipe", "Keterangan"],
    [
        ["status", "VoiceCallStatus", "Status panggilan"],
        ["participants", "VoiceCallParticipantView[]",
         "Daftar peserta lain + indikator berbicara"],
        ["micEnabled", "boolean", "Status mikrofon lokal"],
        ["error", "string | null", "Pesan kesalahan terakhir"],
        ["join(accessToken)", "Promise<void>", "Bergabung ke panggilan"],
        ["leave()", "void", "Keluar dari panggilan"],
        ["toggleMic()", "void", "Aktif/bisukan mikrofon"],
    ],
    widths=[3.6, 4.6, 7.8],
)
para("Contoh", bold=True, size=10)
code(
    "'use client';\n"
    "import { useVoiceCall } from '@teleconf/voicecall-sdk';\n\n"
    "function CallUI({ accessToken }: { accessToken: string }) {\n"
    "  const call = useVoiceCall({ serverUrl: 'https://voice.contoh.com' });\n"
    "  return (\n"
    "    <div>\n"
    "      <p>Status: {call.status}</p>\n"
    "      <button onClick={() => call.join(accessToken)}>Gabung</button>\n"
    "      <button onClick={call.toggleMic}>\n"
    "        {call.micEnabled ? 'Bisukan' : 'Aktifkan'}\n"
    "      </button>\n"
    "      <button onClick={call.leave}>Keluar</button>\n"
    "      <ul>\n"
    "        {call.participants.map((p) => (\n"
    "          <li key={p.peerId}>\n"
    "            {p.displayName} {p.speaking ? '(berbicara)' : ''}\n"
    "          </li>\n"
    "        ))}\n"
    "      </ul>\n"
    "    </div>\n"
    "  );\n"
    "}"
)

# --- 5.3 VoiceCallRoom -----------------------------------------------------
h2("5.3. Komponen VoiceCallRoom")
para("Komponen UI panggilan siap pakai — avatar peserta, indikator pembicara, "
     "tombol mute & keluar.")
h3("Props")
table(
    ["Prop", "Tipe", "Wajib", "Keterangan"],
    [
        ["serverUrl", "string", "Ya", "Base URL server Voice Call"],
        ["accessToken", "string", "Ya", "Token dari endpoint /token"],
        ["selfName", "string", "Tidak", "Nama Anda yang ditampilkan lokal"],
        ["autoJoin", "boolean", "Tidak", "Gabung otomatis saat render (default true)"],
        ["onLeave", "() => void", "Tidak", "Dipanggil setelah keluar panggilan"],
        ["className", "string", "Tidak", "Kelas CSS untuk kontainer luar"],
    ],
    widths=[3.0, 2.6, 1.8, 8.6],
)
para("Contoh", bold=True, size=10)
code(
    "'use client';\n"
    "import { VoiceCallRoom } from '@teleconf/voicecall-sdk';\n\n"
    "<VoiceCallRoom\n"
    "  serverUrl=\"https://voice.contoh.com\"\n"
    "  accessToken={accessToken}\n"
    "  selfName=\"Budi\"\n"
    "  onLeave={() => router.back()}\n"
    "/>"
)

doc.add_page_break()

# ===========================================================================
# 6. TIPE DATA
# ===========================================================================
h1("6. Definisi Tipe Data")

h2("6.1. VoiceCallStatus")
code("type VoiceCallStatus =\n"
     "  | 'idle'        // belum melakukan apa-apa\n"
     "  | 'connecting'  // sedang menyambung\n"
     "  | 'connected'   // panggilan aktif\n"
     "  | 'error'       // terjadi kesalahan\n"
     "  | 'ended';      // panggilan berakhir")

h2("6.2. VoiceCallParticipant")
para("Dipakai pada event remoteStream.")
code("interface VoiceCallParticipant {\n"
     "  peerId: string;        // id unik peserta terhubung\n"
     "  displayName: string;   // nama peserta\n"
     "  stream: MediaStream;   // audio peserta — pasang ke <audio>\n"
     "}")

h2("6.3. VoiceCallParticipantView")
para("Dipakai pada hook useVoiceCall (daftar participants).")
code("interface VoiceCallParticipantView {\n"
     "  peerId: string;\n"
     "  displayName: string;\n"
     "  speaking: boolean;     // true bila sedang menjadi pembicara aktif\n"
     "}")

# ===========================================================================
# 7. CONTOH END-TO-END
# ===========================================================================
h1("7. Contoh Integrasi End-to-End")
para("Contoh lengkap aplikasi Next.js: server menerbitkan token, klien "
     "bergabung ke panggilan.")

h3("Berkas .env aplikasi Anda")
code("VOICECALL_SERVER=https://voice.contoh.com\n"
     "VOICECALL_API_KEY=vc_3f9a1b2c4d5e6f7a8b9c0d1e2f3a4b5c6d7e8f9a")

h3("Server — app/api/voicecall/route.ts")
code(
    "const BASE = process.env.VOICECALL_SERVER!;\n"
    "const KEY  = process.env.VOICECALL_API_KEY!;\n\n"
    "export async function POST(req: Request) {\n"
    "  const { displayName, callId } = await req.json();\n\n"
    "  // buat call baru bila belum ada (panggilan grup memakai callId sama)\n"
    "  let id = callId;\n"
    "  if (!id) {\n"
    "    const c = await fetch(`${BASE}/api/v1/calls`, {\n"
    "      method: 'POST',\n"
    "      headers: { 'X-Api-Key': KEY },\n"
    "    }).then((r) => r.json());\n"
    "    id = c.callId;\n"
    "  }\n\n"
    "  const t = await fetch(`${BASE}/api/v1/calls/${id}/token`, {\n"
    "    method: 'POST',\n"
    "    headers: { 'X-Api-Key': KEY, 'Content-Type': 'application/json' },\n"
    "    body: JSON.stringify({ displayName }),\n"
    "  }).then((r) => r.json());\n\n"
    "  return Response.json({ callId: id, accessToken: t.accessToken });\n"
    "}"
)

h3("Klien — app/call/page.tsx")
code(
    "'use client';\n"
    "import { useState } from 'react';\n"
    "import { VoiceCallRoom } from '@teleconf/voicecall-sdk';\n\n"
    "export default function CallPage() {\n"
    "  const [token, setToken] = useState<string | null>(null);\n\n"
    "  async function start() {\n"
    "    const res = await fetch('/api/voicecall', {\n"
    "      method: 'POST',\n"
    "      headers: { 'Content-Type': 'application/json' },\n"
    "      body: JSON.stringify({ displayName: 'Budi' }),\n"
    "    }).then((r) => r.json());\n"
    "    setToken(res.accessToken);\n"
    "  }\n\n"
    "  if (!token) return <button onClick={start}>Mulai Panggilan</button>;\n"
    "  return (\n"
    "    <VoiceCallRoom\n"
    "      serverUrl=\"https://voice.contoh.com\"\n"
    "      accessToken={token}\n"
    "    />\n"
    "  );\n"
    "}"
)
note("Tombol Mulai memenuhi syarat 'user gesture' sehingga browser "
     "mengizinkan akses mikrofon dan pemutaran audio.")

# ===========================================================================
# 8. DAFTAR KODE STATUS
# ===========================================================================
h1("8. Daftar Kode Status & Error")
table(
    ["Konteks", "Kode / Pesan", "Penjelasan & Tindakan"],
    [
        ["REST", "401 Missing X-Api-Key header",
         "Header API Key belum disertakan"],
        ["REST", "401 Invalid or inactive API key",
         "API Key salah atau dinonaktifkan"],
        ["REST", "400 displayName is required",
         "Body /token tidak memuat displayName"],
        ["REST", "404 Call not found",
         "callId salah atau bukan milik aplikasi Anda"],
        ["REST", "409 Call has already ended",
         "Panggilan sudah diakhiri — buat call baru"],
        ["SDK", "Error: Invalid token",
         "Access token kedaluwarsa — terbitkan token baru"],
        ["SDK", "Error: Call has ended",
         "Panggilan sudah diakhiri saat mencoba bergabung"],
        ["SDK", "status = 'error'",
         "Periksa event 'error' untuk detail penyebab"],
    ],
    widths=[2.2, 5.6, 8.2],
)

doc.add_paragraph()
para("Referensi penggunaan langkah demi langkah tersedia pada dokumen "
     "'Manual Book - SDK Voice Call'.", italic=True, size=9.5, color=GREY)
doc.add_paragraph()
para("--- Akhir Dokumen ---", italic=True, size=9, color=GREY,
     align=WD_ALIGN_PARAGRAPH.CENTER)

out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   "API-Documentation-SDK-VoiceCall.docx")
doc.save(out)
print("Dokumen tersimpan:", out)
