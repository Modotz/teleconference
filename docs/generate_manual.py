# -*- coding: utf-8 -*-
"""
Generator: Manual Book — SDK Voice Call.
Jalankan:  python generate_manual.py
Output  :  Manual-Book-SDK-VoiceCall.docx
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0x25, 0x63, 0xEB)
GREY = RGBColor(0x55, 0x55, 0x55)
HEADER_FILL = "1F3A5F"
ALT_FILL = "EEF2F7"
CODE_FILL = "F2F3F5"

doc = Document()
s = doc.sections[0]
s.left_margin = s.right_margin = Cm(2.4)
s.top_margin = s.bottom_margin = Cm(2.2)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def h1(text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.color.rgb = NAVY
    r.font.size = Pt(16)
    return p


def h2(text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    r.font.color.rgb = ACCENT
    r.font.size = Pt(13)
    return p


def h3(text):
    p = doc.add_heading(level=3)
    r = p.add_run(text)
    r.font.color.rgb = NAVY
    r.font.size = Pt(11.5)
    return p


def para(text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if align is not None:
        p.alignment = align
    return p


def bullet(text, prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if prefix:
        r = p.add_run(prefix)
        r.bold = True
    p.add_run(text)
    return p


def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def code(text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), CODE_FILL)
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    return p


def note(text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFF7E0")
    pPr.append(shd)
    r = p.add_run("  Catatan: " + text)
    r.font.size = Pt(9.5)
    r.italic = True
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(htext)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        shade_cell(c, HEADER_FILL)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            r = cells[ci].paragraphs[0].add_run(str(val))
            r.font.size = Pt(9.5)
            if ri % 2 == 1:
                shade_cell(cells[ci], ALT_FILL)
    if widths:
        for row in t.rows:
            for ci, w in enumerate(widths):
                row.cells[ci].width = Cm(w)
    return t


# ===========================================================================
# COVER
# ===========================================================================
for _ in range(4):
    doc.add_paragraph()
para("MANUAL BOOK", bold=True, size=15, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER)
para("SDK Voice Call", bold=True, size=30, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
para("Buku Panduan Integrasi untuk Pengembang Aplikasi", size=13, color=GREY,
     align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(3):
    doc.add_paragraph()
table(
    ["Keterangan", "Detail"],
    [
        ["Produk", "SDK Voice Call (@teleconf/voicecall-sdk)"],
        ["Versi SDK", "0.1.0"],
        ["Versi Dokumen", "1.0"],
        ["Tanggal", "Mei 2026"],
        ["Untuk", "Tim Pengembang / Integrator"],
    ],
    widths=[5, 11],
)
doc.add_page_break()

# ===========================================================================
# DAFTAR ISI (manual)
# ===========================================================================
h1("Daftar Isi")
for item in [
    "1. Pendahuluan",
    "2. Konsep Dasar",
    "3. Arsitektur & Alur Kerja",
    "4. Prasyarat",
    "5. Instalasi & Konfigurasi Server",
    "6. Membuat API Key",
    "7. Instalasi SDK pada Aplikasi Klien",
    "8. Panduan Integrasi Langkah demi Langkah",
    "9. Pilihan Cara Pakai SDK",
    "10. Skenario Penggunaan",
    "11. Praktik Keamanan Terbaik",
    "12. Troubleshooting",
    "13. FAQ",
    "14. Glosarium",
    "15. Dukungan Teknis",
]:
    p = doc.add_paragraph()
    p.add_run(item).font.size = Pt(11)
doc.add_page_break()

# ===========================================================================
# 1. PENDAHULUAN
# ===========================================================================
h1("1. Pendahuluan")
para(
    "SDK Voice Call adalah pustaka (library) siap pakai untuk menanamkan fitur "
    "panggilan suara real-time ke dalam aplikasi web, khususnya aplikasi "
    "berbasis Next.js / React. Dengan SDK ini, aplikasi Anda dapat melakukan "
    "panggilan suara 1-lawan-1 maupun grup tanpa perlu membangun infrastruktur "
    "WebRTC dari nol."
)
h2("1.1. Untuk Siapa Dokumen Ini")
bullet("Developer yang akan mengintegrasikan SDK ke dalam aplikasi.")
bullet("Administrator sistem yang memasang (deploy) layanan backend.")
h2("1.2. Apa yang Akan Anda Pelajari")
bullet("Cara memasang layanan backend Voice Call.")
bullet("Cara menerbitkan API Key dan access token.")
bullet("Cara memasang dan memakai SDK pada aplikasi Next.js / React.")
bullet("Praktik keamanan dan penyelesaian masalah umum.")

# ===========================================================================
# 2. KONSEP DASAR
# ===========================================================================
h1("2. Konsep Dasar")
para("Sebelum memulai, pahami empat istilah inti berikut:")
table(
    ["Istilah", "Penjelasan"],
    [
        ["Call (Panggilan)",
         "Sebuah sesi panggilan suara. Memiliki callId unik. Satu call bisa "
         "diisi dua peserta (1-on-1) atau banyak peserta (grup)."],
        ["API Key",
         "Kredensial rahasia milik aplikasi Anda. Dipakai HANYA di sisi server "
         "untuk memanggil REST API. Tidak boleh dibawa ke browser."],
        ["Access Token",
         "Token berumur pendek (default 15 menit) untuk SATU peserta pada SATU "
         "panggilan. Token inilah yang dibawa ke browser dan dipakai SDK."],
        ["Participant (Peserta)",
         "Pengguna yang bergabung ke panggilan. Bersifat tamu - identitas "
         "(nama) berasal dari access token, tidak perlu membuat akun."],
    ],
    widths=[4, 12],
)

# ===========================================================================
# 3. ARSITEKTUR & ALUR KERJA
# ===========================================================================
h1("3. Arsitektur & Alur Kerja")
para("Integrasi melibatkan tiga pihak: server backend SDK, server aplikasi "
     "Anda, dan browser pengguna.")
code(
    "  SERVER APLIKASI ANDA            SERVER VOICE CALL\n"
    "  --------------------            -----------------\n"
    "  (1) minta buat call  --X-Api-Key-->  POST /api/v1/calls       -> callId\n"
    "  (2) minta token      --X-Api-Key-->  POST /calls/:id/token    -> token\n"
    "                                        \n"
    "  BROWSER PENGGUNA\n"
    "  ----------------\n"
    "  (3) terima token dari server aplikasi Anda\n"
    "  (4) SDK.join(token)  --WSS/WebRTC-->  Server Voice Call (Mediasoup)\n"
    "  (5) audio mengalir antar peserta"
)
para("Prinsip penting:", bold=True)
bullet("API Key dipakai di server Anda — JANGAN PERNAH dikirim ke browser.")
bullet("Access token boleh dibawa ke browser; umurnya pendek dan terbatas "
       "pada satu panggilan.")
bullet("Media audio mengalir langsung antara browser dan server Mediasoup; "
       "server hanya meneruskan (SFU).")

# ===========================================================================
# 4. PRASYARAT
# ===========================================================================
h1("4. Prasyarat")
h2("4.1. Sisi Server (Backend Voice Call)")
bullet("Node.js versi 20 atau lebih baru.")
bullet("PostgreSQL versi 14 atau lebih baru.")
bullet("Server dengan IP publik dan domain (untuk produksi).")
bullet("Sertifikat TLS (HTTPS) — wajib untuk akses non-localhost.")
bullet("Port terbuka: 80, 443, dan rentang 40000-40100 (UDP + TCP).")
h2("4.2. Sisi Aplikasi Klien")
bullet("Aplikasi berbasis Next.js atau React (versi 18+).")
bullet("Kemampuan menjalankan kode di sisi server (route handler / API route) "
       "untuk menyimpan API Key dengan aman.")
h2("4.3. Sisi Pengguna Akhir")
bullet("Browser modern: Chrome, Edge, Firefox, atau Safari versi terbaru.")
bullet("Izin akses mikrofon.")
bullet("Konteks aman: halaman harus diakses via https:// atau http://localhost.")

doc.add_page_break()

# ===========================================================================
# 5. INSTALASI & KONFIGURASI SERVER
# ===========================================================================
h1("5. Instalasi & Konfigurasi Server")
para("Langkah ini dilakukan oleh administrator untuk menjalankan layanan "
     "backend Voice Call.")
h2("5.1. Menyiapkan Berkas Konfigurasi (.env)")
para("Salin berkas contoh lalu sesuaikan nilainya:")
code("cd backend\n"
     "Copy-Item .env.example .env")
para("Nilai penting pada .env:")
table(
    ["Variabel", "Keterangan"],
    [
        ["DATABASE_URL", "Koneksi PostgreSQL"],
        ["JWT_SECRET", "Kunci rahasia token login pengguna"],
        ["CALL_TOKEN_SECRET", "Kunci rahasia penanda-tanganan access token"],
        ["CALL_TOKEN_TTL", "Masa berlaku access token, mis. 15m"],
        ["MEDIASOUP_ANNOUNCED_IP", "IP publik server (untuk produksi)"],
        ["MEDIASOUP_MIN_PORT / MAX_PORT", "Rentang port media (40000-40100)"],
    ],
    widths=[5.5, 10.5],
)
h2("5.2. Memasang Dependensi & Basis Data")
code("npm install\n"
     "npx prisma migrate dev --name init")
h2("5.3. Menjalankan Server")
code("# Mode HTTPS (produksi / LAN)\n"
     "npm run dev:https\n\n"
     "# atau mode HTTP (khusus localhost)\n"
     "npm run dev")
note("Untuk produksi gunakan process manager (mis. PM2 atau systemd) dan "
     "reverse proxy Nginx dengan sertifikat TLS.")

# ===========================================================================
# 6. MEMBUAT API KEY
# ===========================================================================
h1("6. Membuat API Key")
para("Setiap aplikasi pihak ketiga membutuhkan satu API Key. Buat dengan "
     "perintah berikut di folder backend:")
code('node scripts/create-api-client.js "Nama Aplikasi Partner"')
para("Perintah ini mencetak API Key SATU KALI. Contoh keluaran:")
code("API client created\n"
     "  name : Nama Aplikasi Partner\n"
     "  id   : 7c1e...\n\n"
     "API KEY (shown once - store it now):\n"
     "  vc_3f9a1b2c4d5e6f7a8b9c0d1e2f3a4b5c6d7e8f9a")
note("Simpan API Key segera di tempat aman (mis. variabel lingkungan aplikasi "
     "partner). Kunci tidak dapat ditampilkan ulang — hanya hash-nya disimpan.")

# ===========================================================================
# 7. INSTALASI SDK
# ===========================================================================
h1("7. Instalasi SDK pada Aplikasi Klien")
para("SDK didistribusikan sebagai paket npm bernama @teleconf/voicecall-sdk.")
h2("7.1. Instalasi dari Registry")
code("npm install @teleconf/voicecall-sdk")
h2("7.2. Instalasi dari Berkas Lokal (distribusi privat)")
code("# di folder sdk:\n"
     "npm install\n"
     "npm run build\n"
     "npm pack            # menghasilkan teleconf-voicecall-sdk-0.1.0.tgz\n\n"
     "# di aplikasi partner:\n"
     "npm install ../path/teleconf-voicecall-sdk-0.1.0.tgz")

doc.add_page_break()

# ===========================================================================
# 8. PANDUAN INTEGRASI LANGKAH DEMI LANGKAH
# ===========================================================================
h1("8. Panduan Integrasi Langkah demi Langkah")

h3("Langkah 1 — Simpan API Key di Server Aplikasi Anda")
para("Tambahkan API Key ke berkas .env aplikasi Anda (jangan di kode klien):")
code("VOICECALL_API_KEY=vc_3f9a1b2c4d5e6f7a8b9c0d1e2f3a4b5c6d7e8f9a\n"
     "VOICECALL_SERVER=https://voice.contoh.com")

h3("Langkah 2 — Buat Endpoint Server untuk Menerbitkan Token")
para("Buat route handler yang membuat call dan menerbitkan token. API Key "
     "tetap aman di server.")
code(
    "// app/api/voicecall/route.ts  (Next.js App Router)\n"
    "const BASE = process.env.VOICECALL_SERVER!;\n"
    "const KEY  = process.env.VOICECALL_API_KEY!;\n\n"
    "export async function POST(req: Request) {\n"
    "  const { displayName } = await req.json();\n\n"
    "  // 1. Buat call\n"
    "  const call = await fetch(`${BASE}/api/v1/calls`, {\n"
    "    method: 'POST',\n"
    "    headers: { 'X-Api-Key': KEY },\n"
    "  }).then((r) => r.json());\n\n"
    "  // 2. Terbitkan token untuk peserta ini\n"
    "  const token = await fetch(\n"
    "    `${BASE}/api/v1/calls/${call.callId}/token`,\n"
    "    {\n"
    "      method: 'POST',\n"
    "      headers: { 'X-Api-Key': KEY, 'Content-Type': 'application/json' },\n"
    "      body: JSON.stringify({ displayName }),\n"
    "    }\n"
    "  ).then((r) => r.json());\n\n"
    "  return Response.json({\n"
    "    callId: call.callId,\n"
    "    accessToken: token.accessToken,\n"
    "  });\n"
    "}"
)
note("Untuk panggilan grup: buat call SATU KALI, lalu panggil endpoint token "
     "berkali-kali dengan callId yang SAMA untuk tiap peserta.")

h3("Langkah 3 — Tampilkan Panggilan di Sisi Klien")
para("Gunakan komponen siap pakai <VoiceCallRoom>:")
code(
    "'use client';\n"
    "import { VoiceCallRoom } from '@teleconf/voicecall-sdk';\n\n"
    "export default function CallPage({ accessToken }: { accessToken: string }) {\n"
    "  return (\n"
    "    <VoiceCallRoom\n"
    "      serverUrl=\"https://voice.contoh.com\"\n"
    "      accessToken={accessToken}\n"
    "      selfName=\"Budi\"\n"
    "      onLeave={() => console.log('keluar dari panggilan')}\n"
    "    />\n"
    "  );\n"
    "}"
)

# ===========================================================================
# 9. PILIHAN CARA PAKAI SDK
# ===========================================================================
h1("9. Pilihan Cara Pakai SDK")
para("SDK menyediakan tiga tingkat penggunaan, dari paling mudah hingga "
     "paling fleksibel.")

h2("9.1. Komponen <VoiceCallRoom> (paling mudah)")
para("UI panggilan siap pakai: avatar peserta, indikator pembicara, tombol "
     "mute dan keluar. Cukup berikan serverUrl dan accessToken.")

h2("9.2. Hook useVoiceCall (UI bebas)")
para("Logika panggilan tanpa UI — Anda merancang tampilan sendiri.")
code(
    "'use client';\n"
    "import { useVoiceCall } from '@teleconf/voicecall-sdk';\n\n"
    "function MyCall({ accessToken }: { accessToken: string }) {\n"
    "  const call = useVoiceCall({ serverUrl: 'https://voice.contoh.com' });\n"
    "  return (\n"
    "    <div>\n"
    "      <p>Status: {call.status}</p>\n"
    "      <button onClick={() => call.join(accessToken)}>Gabung</button>\n"
    "      <button onClick={call.toggleMic}>\n"
    "        {call.micEnabled ? 'Bisukan' : 'Aktifkan Mic'}\n"
    "      </button>\n"
    "      <button onClick={call.leave}>Keluar</button>\n"
    "    </div>\n"
    "  );\n"
    "}"
)
note("Audio peserta lain diputar otomatis oleh SDK. Anda hanya merancang UI.")

h2("9.3. Kelas VoiceCallClient (kontrol penuh)")
para("Untuk integrasi tingkat lanjut tanpa React.")
code(
    "import { VoiceCallClient } from '@teleconf/voicecall-sdk';\n\n"
    "const client = new VoiceCallClient({\n"
    "  serverUrl: 'https://voice.contoh.com',\n"
    "});\n"
    "client.on('remoteStream', ({ stream }) => {\n"
    "  const a = new Audio();\n"
    "  a.srcObject = stream;\n"
    "  a.play();\n"
    "});\n"
    "await client.join(accessToken);\n"
    "client.setMicEnabled(false);   // bisukan\n"
    "client.leave();"
)

doc.add_page_break()

# ===========================================================================
# 10. SKENARIO PENGGUNAAN
# ===========================================================================
h1("10. Skenario Penggunaan")
h2("10.1. Panggilan 1-lawan-1")
numbered("Server aplikasi membuat satu call.")
numbered("Server menerbitkan token untuk penelepon dan penerima (2 token).")
numbered("Kedua pihak membuka halaman panggilan dengan token masing-masing.")
numbered("SDK menyambungkan keduanya; audio mengalir dua arah.")

h2("10.2. Panggilan Grup")
numbered("Server aplikasi membuat satu call.")
numbered("Untuk setiap anggota grup, server menerbitkan token dengan callId "
         "yang SAMA.")
numbered("Semua anggota bergabung memakai token masing-masing.")
numbered("SDK secara otomatis menyambungkan semua peserta dalam satu ruang.")

h2("10.3. Mengakhiri Panggilan")
bullet("Peserta keluar sendiri: panggil leave() atau tombol Keluar.")
bullet("Mengakhiri seluruh panggilan: server memanggil POST /calls/:id/end.")

# ===========================================================================
# 11. KEAMANAN
# ===========================================================================
h1("11. Praktik Keamanan Terbaik")
bullet("JANGAN menaruh API Key di kode frontend / browser. Hanya di server.",
       "API Key. ")
bullet("Selalu gunakan HTTPS untuk produksi. Mikrofon hanya aktif pada "
       "konteks aman.", "HTTPS. ")
bullet("Terbitkan access token sesaat sebelum dibutuhkan; umurnya pendek.",
       "Token. ")
bullet("Terbitkan token hanya setelah memverifikasi bahwa pengguna berhak "
       "ikut panggilan tersebut.", "Otorisasi. ")
bullet("Simpan API Key di variabel lingkungan, bukan di repositori kode.",
       "Penyimpanan. ")
bullet("Ganti API Key bila dicurigai bocor (buat ApiClient baru).",
       "Rotasi. ")

# ===========================================================================
# 12. TROUBLESHOOTING
# ===========================================================================
h1("12. Troubleshooting")
table(
    ["Gejala", "Kemungkinan Penyebab", "Solusi"],
    [
        ["SDK gagal join, error 'Invalid token'",
         "Access token kedaluwarsa atau salah",
         "Terbitkan token baru tepat sebelum join"],
        ["Mikrofon tidak aktif / diblokir",
         "Halaman bukan konteks aman, atau izin ditolak",
         "Akses via https:// dan izinkan mikrofon di browser"],
        ["Tidak terdengar suara peserta lain",
         "Autoplay diblokir browser",
         "Panggil join() dari aksi klik pengguna"],
        ["Gagal terhubung / connect_error",
         "Port media tertutup atau IP announce salah",
         "Buka port 40000-40100 UDP/TCP, set MEDIASOUP_ANNOUNCED_IP"],
        ["Error 401 saat memanggil REST API",
         "API Key salah / tidak aktif",
         "Periksa header X-Api-Key; buat ulang bila perlu"],
        ["Error 409 saat menerbitkan token",
         "Call sudah diakhiri",
         "Buat call baru terlebih dahulu"],
    ],
    widths=[4.8, 5.4, 5.8],
)

# ===========================================================================
# 13. FAQ
# ===========================================================================
h1("13. FAQ")


def faq(q, a):
    para("T: " + q, bold=True, size=10.5)
    para("J: " + a, size=10.5)
    doc.add_paragraph()


faq("Apakah peserta perlu membuat akun?",
    "Tidak. Peserta adalah tamu; identitas berasal dari access token yang "
    "diterbitkan server aplikasi Anda.")
faq("Berapa lama access token berlaku?",
    "Secara default 15 menit (dapat diatur lewat CALL_TOKEN_TTL). Ini adalah "
    "jendela untuk JOIN; setelah tersambung, sesi tetap berjalan.")
faq("Berapa peserta maksimum dalam satu panggilan grup?",
    "Secara teknis besar; untuk kualitas terbaik disarankan hingga belasan "
    "peserta per panggilan. Kapasitas total bergantung spesifikasi server.")
faq("Apakah SDK mendukung video?",
    "Tidak. SDK ini khusus panggilan suara. Fitur video tersedia pada paket "
    "Video Meeting terpisah.")
faq("Apakah bisa dipakai di aplikasi mobile?",
    "Bisa, melalui web view / browser mobile. SDK berbasis web (WebRTC).")

# ===========================================================================
# 14. GLOSARIUM
# ===========================================================================
h1("14. Glosarium")
table(
    ["Istilah", "Arti"],
    [
        ["SDK", "Software Development Kit — pustaka siap pakai untuk developer"],
        ["WebRTC", "Standar web untuk komunikasi audio/video real-time"],
        ["SFU", "Selective Forwarding Unit — server yang meneruskan media"],
        ["Signaling", "Pertukaran pesan untuk membangun koneksi panggilan"],
        ["Access Token", "Token berumur pendek untuk satu peserta/panggilan"],
        ["API Key", "Kredensial rahasia aplikasi, dipakai di sisi server"],
        ["Codec Opus", "Format kompresi audio berkualitas tinggi"],
    ],
    widths=[4, 12],
)

# ===========================================================================
# 15. DUKUNGAN
# ===========================================================================
h1("15. Dukungan Teknis")
para("Untuk pertanyaan teknis, kendala integrasi, atau permintaan dukungan, "
     "hubungi tim kami melalui kanal yang disepakati pada kontrak layanan.")
table(
    ["Kanal", "Detail"],
    [
        ["Email Dukungan", "[email dukungan Anda]"],
        ["Jam Layanan", "[jam kerja layanan]"],
        ["Dokumentasi API", "https://[server]/api/v1/docs"],
        ["Referensi API", "Lihat dokumen 'API Documentation - SDK Voice Call'"],
    ],
    widths=[5, 11],
)
doc.add_paragraph()
para("--- Akhir Dokumen ---", italic=True, size=9, color=GREY,
     align=WD_ALIGN_PARAGRAPH.CENTER)

out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   "Manual-Book-SDK-VoiceCall.docx")
doc.save(out)
print("Dokumen tersimpan:", out)
